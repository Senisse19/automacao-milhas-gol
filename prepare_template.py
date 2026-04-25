from docx import Document
import os

# Caminho do documento original
original_doc = "TERMO DE QUITACAO MILHAS.docx"

if not os.path.exists(original_doc):
    print(f"Erro: Arquivo '{original_doc}' não encontrado.")
    exit(1)

doc = Document(original_doc)

replacements = {
    "SISBSU": "{CODIGO_LOCALIZADOR}",
    "POAG360019": "{RELATORIO_DANO}",
    "PRISCILA BRAGA PEREIRA": "{NOME_CLIENTE}",
    "10095208739": "{CPF_CLIENTE}",
    "AV SAO JOSE MARIA ESCRIVA 560 BLOCO 11611 CEP 22753-200": "{ENDERECO} {BAIRRO} CEP {CEP}",
    "366002033": "{NUMERO_SMILES}",
    "25.000": "{QUANTIDADE_MILHAS}",
    "PORTO ALEGRE, 19 DE ABRIL DE 2026": "{DATA_EXTENSO}",
    "PORTO ALEGRE": "{CIDADE_UF}",
}

def replace_in_paragraphs(paragraphs, replacements):
    for paragraph in paragraphs:
        for old, new in replacements.items():
            if old in paragraph.text:
                for run in paragraph.runs:
                    if old in run.text:
                        run.text = run.text.replace(old, new)

replace_in_paragraphs(doc.paragraphs, replacements)

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            replace_in_paragraphs(cell.paragraphs, replacements)

# Salvar na pasta public do Next.js
output_dir = "public"
if not os.path.exists(output_dir):
    os.makedirs(output_dir)

output_path = os.path.join(output_dir, "template.docx")
doc.save(output_path)
print(f"Template salvo em {output_path}")
print("REVISE O ARQUIVO NO WORD ANTES DE FAZER O DEPLOY.")
